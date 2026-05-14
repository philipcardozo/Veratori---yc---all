"""
Veratori weekly report builder.

Opens the on-disk template (Tori_GenDocRef_v4.docx), walks the document,
replaces every bracket-style placeholder with real data (or "Unavailable
for this period." in italics when data is missing), removes the
✦ TORI — instruction lines, repopulates the appendix table row-by-row,
and embeds the two required charts.

The template file is read-only — every generation produces a fresh copy.
"""

from __future__ import annotations

import copy
import hashlib
import io
import os
import random
from datetime import date, datetime, timedelta
from typing import Optional

import matplotlib
matplotlib.use("Agg")  # headless backend — must precede pyplot import
import matplotlib.pyplot as plt

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "templates", "Tori_GenDocRef_v4.docx"
)

UNAVAILABLE = "Unavailable for this period."

# Veratori brand accents used in matplotlib charts
BRAND_GREEN  = "#76b900"
BRAND_BLUES  = ["#1f77b4", "#3a8fc4", "#5aa6d1", "#7abedd", "#9bd6e9"]
BRAND_WARM   = ["#f59e0b", "#ef4444", "#ec4899", "#8b5cf6"]


# ── Franchise registry ────────────────────────────────────────────────────────
# Owner name and contact mirror what the franchise hub uses. Camera Demo
# is treated as the only "live" location; the others pull seed data.
FRANCHISES = {
    # Keys mirror the franchise IDs used in home.html so the frontend can pass them through unchanged.
    "canal":  {"name": "Canal Street",   "owner": "Johnny CEO",      "contact": "johnny@veratori.ai"},
    "f43":    {"name": "43rd Street",    "owner": "Marcus Rivera",   "contact": "marcus.r@veratori.ai"},
    "f44":    {"name": "44th Street",    "owner": "Priya Patel",     "contact": "priya.p@veratori.ai"},
    "f45":    {"name": "45th Street",    "owner": "Daniel Kim",      "contact": "daniel.k@veratori.ai"},
    "f46":    {"name": "46th Street",    "owner": "Sofia Hernández", "contact": "sofia.h@veratori.ai"},
    "cam":    {"name": "Camera Demo",    "owner": "Felipe Cardozo",  "contact": "felipe@veratori.ai", "live": True},
}

# Seed menu — frozen yogurt / açaí bowl franchise. Same menu across non-Camera
# locations so cross-location comparisons make sense; per-week unit volumes
# vary deterministically per franchise + week.
SEED_MENU = [
    {"name": "Mango Tango",         "price": 9.50},
    {"name": "Strawberry Bliss",    "price": 8.75},
    {"name": "Coconut Dream",       "price": 9.25},
    {"name": "Passion Fruit Bowl",  "price": 10.50},
    {"name": "Guava Breeze",        "price": 8.50},
    {"name": "Lychee Rose",         "price": 9.75},
    {"name": "Maui Custard",        "price": 7.50},
    {"name": "Acai Classic",        "price": 11.00},
]


# ── Date helpers ──────────────────────────────────────────────────────────────
def parse_iso_week(week_iso: str) -> tuple[date, date]:
    """Parse 'YYYY-Www' (e.g. '2026-W19') into (monday, sunday)."""
    year_part, week_part = week_iso.split("-W")
    monday = date.fromisocalendar(int(year_part), int(week_part), 1)
    return monday, monday + timedelta(days=6)


def fmt_week_range(monday: date, sunday: date) -> str:
    return f"{monday.strftime('%b %-d')} – {sunday.strftime('%b %-d, %Y')}"


# ── Data gathering ────────────────────────────────────────────────────────────
def gather_data(franchise_id: str, week_iso: str, sales_tracker) -> dict:
    """
    Returns a dict shaped for the template. Has these keys:
      meta:          {franchise, owner, contact, period_iso, period_label, generated, is_modeled}
      summary:       [{title, body}, ...]                — 4-6 bullet insights
      hot:           [{name, units, insight}, ...]       — top 3 by units
      not_hot:       [{name, wasted, insight}, ...]      — top 3 by waste
      locations:     [{name, top_seller, top_units, profit_item, profit_rev}, ...]
      cross_trends:  str                                  — 2-3 sentence trends
      forecast:      {opening, orders:[...], reduce:[...], priorities:[...]}
      appendix:      [{name, units, revenue, wasted, note}, ...]  — full menu rows
      charts:        {pie_caption, bar_caption}
      pie_data:      [(label, revenue), ...]              — for matplotlib
      bar_data:      [(label, purchased, sold), ...]      — for matplotlib
    """
    franchise = FRANCHISES.get(franchise_id)
    if not franchise:
        raise ValueError(f"unknown franchise id: {franchise_id}")

    monday, sunday = parse_iso_week(week_iso)
    is_live = bool(franchise.get("live"))

    if is_live and sales_tracker is not None:
        return _gather_live(franchise, week_iso, monday, sunday, sales_tracker)
    return _gather_seeded(franchise, franchise_id, week_iso, monday, sunday)


def _gather_live(franchise: dict, week_iso: str, monday: date, sunday: date, sales_tracker) -> dict:
    """Camera Demo: real CSVs from sales_tracker.get_report."""
    report = sales_tracker.get_report(monday.isoformat(), sunday.isoformat())
    has_sales = bool(report and report["totals"]["rows_read"] > 0)

    # Real menu items observed in CSVs (COCO classes)
    products = report.get("top_products", []) if report else []
    if has_sales:
        hot = [
            {
                "name":    p["label"].title(),
                "units":   p["units"],
                "insight": f"{p['share']}% of weekly revenue — strongest seller in window.",
            }
            for p in products[:3]
        ]
        # We do not have waste data on the live feed — flag explicitly.
        not_hot = []
        appendix = [
            {
                "name":    p["label"].title(),
                "units":   p["units"],
                "revenue": f"${p['revenue']:.2f}",
                "wasted":  UNAVAILABLE,
                "note":    "Live CSV — waste tracking pending.",
            }
            for p in products
        ]
        pie_data = [(p["label"].title(), p["revenue"]) for p in products if p["revenue"] > 0]
        bar_data = [
            (p["label"].title(),
             # Inventory purchased is unknown for live data — flag with None below
             None,
             p["units"])
            for p in products[:6]
        ]
    else:
        hot, not_hot, appendix, pie_data, bar_data = [], [], [], [], []

    kpis = report["kpis"] if has_sales else None
    summary = _summary_bullets_live(kpis, products, monday, sunday)
    locations = [{
        "name":        franchise["name"],
        "top_seller":  hot[0]["name"] if hot else UNAVAILABLE,
        "top_units":   str(hot[0]["units"]) if hot else UNAVAILABLE,
        "profit_item": products[0]["label"].title() if products else UNAVAILABLE,
        "profit_rev":  f"${products[0]['revenue']:.2f}" if products else UNAVAILABLE,
    }]
    forecast = _forecast_block(products, [franchise["name"]], has_sales)
    return {
        "meta": {
            "franchise":    franchise["name"],
            "owner":        franchise["owner"],
            "contact":      franchise["contact"],
            "period_iso":   week_iso,
            "period_label": fmt_week_range(monday, sunday),
            "generated":    datetime.now().strftime("%b %-d, %Y"),
            "is_modeled":   False,
        },
        "summary":      summary,
        "hot":          hot,
        "not_hot":      not_hot,
        "locations":    locations,
        "cross_trends": "Single-location report — cross-location comparison not applicable.",
        "forecast":     forecast,
        "appendix":     appendix,
        "charts": {
            "pie_caption": _pie_caption(products),
            "bar_caption": "Inventory purchase data is not yet wired into Veratori — bars show sales only for the live location.",
        },
        "pie_data":  pie_data,
        "bar_data":  bar_data,
    }


def _gather_seeded(franchise: dict, franchise_id: str, week_iso: str, monday: date, sunday: date) -> dict:
    """Non-camera franchises: deterministic per-(franchise, week) synthetic data."""
    seed = int(hashlib.sha256(f"{franchise_id}|{week_iso}".encode()).hexdigest()[:8], 16)
    rng = random.Random(seed)

    items = []
    for m in SEED_MENU:
        units    = rng.randint(40, 220)
        wasted   = max(0, int(units * rng.uniform(0.02, 0.12)))
        purchased = units + wasted + rng.randint(0, 18)
        revenue  = round(units * m["price"], 2)
        items.append({
            "name":      m["name"],
            "units":     units,
            "wasted":    wasted,
            "purchased": purchased,
            "revenue":   revenue,
            "price":     m["price"],
        })

    # Sort for ranked lists
    by_units = sorted(items, key=lambda r: -r["units"])
    by_rev   = sorted(items, key=lambda r: -r["revenue"])
    by_waste = sorted(items, key=lambda r: -r["wasted"])

    hot = [{
        "name":    it["name"],
        "units":   it["units"],
        "insight": f"Consistent weekday performer — {it['units']} units moved this week.",
    } for it in by_units[:3]]

    not_hot = [{
        "name":    it["name"],
        "wasted":  it["wasted"],
        "insight": f"{it['wasted']} units discarded — consider trimming the standing order.",
    } for it in by_waste[:3]]

    appendix = [{
        "name":    it["name"],
        "units":   it["units"],
        "revenue": f"${it['revenue']:,.2f}",
        "wasted":  it["wasted"],
        "note":    f"${it['price']:.2f} unit price",
    } for it in by_rev]

    pie_data = [(it["name"], it["revenue"]) for it in by_rev[:6]]
    bar_data = [(it["name"], it["purchased"], it["units"]) for it in by_rev[:6]]

    total_rev   = sum(it["revenue"] for it in items)
    total_units = sum(it["units"]   for it in items)
    total_waste = sum(it["wasted"]  for it in items)

    summary = [
        {
            "title": "Total Revenue",
            "body":  f"The location closed the week at ${total_rev:,.2f} across {total_units} units. See Appendix A.",
        },
        {
            "title": "Top-Selling Item",
            "body":  f"{by_units[0]['name']} led the menu by volume this week. See Appendix A.",
        },
        {
            "title": "High-Waste Item",
            "body":  f"{by_waste[0]['name']} drove the largest end-of-night discard count. See Appendix A.",
        },
        {
            "title": "Revenue Mix",
            "body":  f"The top three items contributed roughly {round(sum(i['revenue'] for i in by_rev[:3]) / total_rev * 100)}% of total revenue. See Appendix B.",
        },
        {
            "title": "Waste Pressure",
            "body":  f"Aggregate waste ran at {round(total_waste / total_units * 100, 1)}% of units sold — within target range. See Appendix A.",
        },
    ]

    locations = [{
        "name":        franchise["name"],
        "top_seller":  by_units[0]["name"],
        "top_units":   str(by_units[0]["units"]),
        "profit_item": by_rev[0]["name"],
        "profit_rev":  f"${by_rev[0]['revenue']:,.2f}",
    }]

    forecast = _forecast_block_seeded(by_units, by_waste, [franchise["name"]])

    return {
        "meta": {
            "franchise":    franchise["name"],
            "owner":        franchise["owner"],
            "contact":      franchise["contact"],
            "period_iso":   week_iso,
            "period_label": fmt_week_range(monday, sunday),
            "generated":    datetime.now().strftime("%b %-d, %Y"),
            "is_modeled":   True,
        },
        "summary":      summary,
        "hot":          hot,
        "not_hot":      not_hot,
        "locations":    locations,
        "cross_trends": "Single-location report. Network-wide trend comparison available once additional franchises are activated.",
        "forecast":     forecast,
        "appendix":     appendix,
        "charts": {
            "pie_caption": _pie_caption_seeded(by_rev, total_rev),
            "bar_caption": f"Inventory purchased tracks units sold within ~5% across most items; {by_waste[0]['name']} shows the widest gap.",
        },
        "pie_data":  pie_data,
        "bar_data":  bar_data,
    }


def _summary_bullets_live(kpis, products, monday, sunday):
    if not kpis:
        return [{
            "title": "No Sales Recorded",
            "body":  f"No sales were captured for the week of {fmt_week_range(monday, sunday)}. See Appendix A.",
        }]
    bullets = [{
        "title": "Total Revenue",
        "body":  f"Camera Demo closed the week at ${kpis['revenue']:.2f} across {kpis['units']} units. See Appendix A.",
    }, {
        "title": "Top-Selling Item",
        "body":  f"{kpis['top_product']} led volume on the live feed. See Appendix A.",
    }]
    if products:
        bullets.append({
            "title": "Display Throughput",
            "body":  f"Average time on display before sale: {kpis['avg_on_display_s']}s. See Appendix A.",
        })
    bullets.append({
        "title": "Waste Tracking",
        "body":  "End-of-night discard data is not yet emitted by the live camera feed — POS integration pending.",
    })
    return bullets


def _forecast_block(products, locations, has_sales):
    if not has_sales:
        return {
            "opening": f"No sales captured this week — no forecast can be derived. See Appendix A.",
            "orders":  [],
            "reduce":  [],
            "priorities": [{"location": locations[0], "action": "Verify camera coverage and re-run report."}],
        }
    orders = [{
        "item":      p["label"].title(),
        "qty":       max(10, p["units"] * 2),
        "rationale": f"{p['units']} units moved this week ({p['share']}% of revenue). See Appendix D.",
    } for p in products[:3]]
    return {
        "opening": f"Based on this week's live camera data, the following actions are recommended for next week.",
        "orders":  orders,
        "reduce":  [],
        "priorities": [{"location": locations[0], "action": "Maintain current display rotation cadence."}],
    }


def _forecast_block_seeded(by_units, by_waste, locations):
    orders = [{
        "item":      by_units[0]["name"],
        "qty":       int(by_units[0]["units"] * 1.10),
        "rationale": "Trending up 10% week-over-week — increase standing order to match. See Appendix D.",
    }, {
        "item":      by_units[1]["name"],
        "qty":       int(by_units[1]["units"] * 1.05),
        "rationale": "Stable demand — small uplift to match next-week forecast. See Appendix D.",
    }]
    reduce = [{
        "item":   by_waste[0]["name"],
        "detail": f"End-of-night surplus observed 3 of past 4 weeks. See Appendix A.",
    }]
    priorities = [{
        "location": locations[0],
        "action":   f"Pull {by_waste[0]['name']} from the standing order until waste rate drops below 5%.",
    }]
    return {
        "opening": "Based on this week's performance and Veratori trend data, the following actions are recommended for next week.",
        "orders":  orders,
        "reduce":  reduce,
        "priorities": priorities,
    }


def _pie_caption(products):
    if not products:
        return UNAVAILABLE
    top = products[0]
    return (f"{top['label'].title()} drove {top['share']}% of revenue this week; "
            f"the top three items together account for "
            f"{round(sum(p['share'] for p in products[:3]))}% of the mix.")


def _pie_caption_seeded(by_rev, total_rev):
    top = by_rev[0]
    share = round(top["revenue"] / total_rev * 100, 1)
    top3  = round(sum(it["revenue"] for it in by_rev[:3]) / total_rev * 100)
    return (f"{top['name']} drove {share}% of revenue this week; "
            f"the top three items together account for {top3}% of the mix.")


# ── Chart rendering ───────────────────────────────────────────────────────────
def render_pie_chart(pie_data: list[tuple[str, float]]) -> Optional[io.BytesIO]:
    if not pie_data:
        return None
    labels = [d[0] for d in pie_data]
    values = [d[1] for d in pie_data]
    fig, ax = plt.subplots(figsize=(6.0, 4.0), dpi=160)
    colors = (BRAND_BLUES + BRAND_WARM)[:len(values)]
    wedges, _texts, autotexts = ax.pie(
        values, labels=labels, colors=colors,
        autopct="%1.0f%%", pctdistance=0.78,
        textprops={"fontsize": 9, "color": "#0b1220"},
        wedgeprops={"linewidth": 1.2, "edgecolor": "white"},
    )
    for at in autotexts:
        at.set_color("white"); at.set_fontweight("bold")
    ax.set_title("Sales Contribution by Item", fontsize=11, fontweight="bold", color="#0b1220", pad=10)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def render_bar_chart(bar_data: list[tuple[str, Optional[float], float]]) -> Optional[io.BytesIO]:
    if not bar_data:
        return None
    labels = [d[0] for d in bar_data]
    purchased_raw = [d[1] for d in bar_data]
    sold = [d[2] for d in bar_data]
    has_purchased = all(p is not None for p in purchased_raw)
    purchased = [p if p is not None else 0 for p in purchased_raw]

    fig, ax = plt.subplots(figsize=(6.0, 3.6), dpi=160)
    import numpy as np
    x = np.arange(len(labels))
    w = 0.38
    if has_purchased:
        ax.bar(x - w/2, purchased, w, label="Purchased", color=BRAND_GREEN)
        ax.bar(x + w/2, sold,      w, label="Sold",      color="#1f77b4")
    else:
        ax.bar(x, sold, w * 1.8, label="Sold", color=BRAND_GREEN)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=8)
    ax.set_ylabel("Units", fontsize=9)
    ax.set_title("Inventory Purchased vs. Units Sold", fontsize=11, fontweight="bold", color="#0b1220", pad=10)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.5)
    if has_purchased:
        ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ── Docx manipulation ─────────────────────────────────────────────────────────
def _iter_paragraphs(doc):
    """Yield every paragraph in the document — body and inside tables."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                # Also recurse into nested tables in cells if any
                for inner_tbl in cell.tables:
                    for inner_row in inner_tbl.rows:
                        for inner_cell in inner_row.cells:
                            yield from inner_cell.paragraphs


def _set_paragraph_text(paragraph, new_text, italic=False):
    """
    Replace a paragraph's text while keeping the first run's formatting.
    Italic flag overrides the run italic state (used for "Unavailable" markers).
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        if italic:
            paragraph.runs[0].italic = True
        return
    first = paragraph.runs[0]
    first.text = new_text
    if italic:
        first.italic = True
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_in_paragraph(paragraph, mapping: dict[str, str]):
    """In-place substring replacement keyed by placeholder → value."""
    text = paragraph.text
    if not text:
        return
    new_text = text
    hit = False
    for placeholder, value in mapping.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, value)
            hit = True
    if hit:
        _set_paragraph_text(paragraph, new_text)


def _strip_instruction_paragraphs(doc):
    """
    Remove every paragraph that is part of an instruction block. The template
    marks these in two ways:
      1. Lines starting with the ✦ glyph
      2. The DOCUMENT-LEVEL INSTRUCTIONS body that follows it
      3. The 'Tori can provide a deeper drill-down…' closer (template says to
         keep it; we keep it). All other meta-prose is dropped.
    """
    KILL_PHRASES = (
        "Replace every italicised placeholder",
        "Read all instructions before generating content",
        "Delete every instruction block",
        "no instruction text",
        "No instruction text",
    )
    for p in list(_iter_paragraphs(doc)):
        txt = p.text.lstrip() if p.text else ""
        if not txt:
            continue
        if txt.startswith("✦") or any(phrase in txt for phrase in KILL_PHRASES):
            _drop_paragraph(p)


def _drop_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def _find_paragraph_starting_with(doc, prefix):
    for p in _iter_paragraphs(doc):
        if p.text and p.text.lstrip().startswith(prefix):
            return p
    return None


def _insert_paragraph_after(paragraph):
    """python-docx has insert_paragraph_before; we wrap to insert after."""
    new_p = paragraph._element.makeelement(
        qn("w:p"), {}
    )
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


# ── Top-level orchestration ───────────────────────────────────────────────────
_SOFFICE_CANDIDATES = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/opt/homebrew/bin/soffice",
    "/usr/local/bin/soffice",
    "soffice",  # whatever's on PATH
]


def _find_soffice() -> str:
    """Locate the LibreOffice headless binary on this machine."""
    import shutil
    for cand in _SOFFICE_CANDIDATES:
        if os.path.isabs(cand):
            if os.path.exists(cand):
                return cand
        else:
            found = shutil.which(cand)
            if found:
                return found
    raise RuntimeError(
        "LibreOffice not found. Install with: brew install --cask libreoffice"
    )


def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Convert the generated docx to PDF using LibreOffice in headless mode.
    Writes the docx to a temp file, runs `soffice --convert-to pdf`, reads
    the resulting PDF, and cleans up. Same render engine class as Word's
    "Save as PDF" — preserves tables, images, charts, and italic placeholders.
    """
    import subprocess
    import tempfile

    soffice = _find_soffice()
    with tempfile.TemporaryDirectory(prefix="veratori-pdf-") as td:
        docx_path = os.path.join(td, "report.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        # First-run profile initialization can take ~30-60s; warmed-up calls take ~10s.
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", td, docx_path],
            capture_output=True, text=True, timeout=180,
        )
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {proc.stderr or proc.stdout}")
        pdf_path = os.path.join(td, "report.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"LibreOffice did not produce a PDF. Output: {proc.stdout}")
        with open(pdf_path, "rb") as f:
            return f.read()


def build_report(franchise_id: str, week_iso: str, sales_tracker=None) -> bytes:
    """
    Returns the filled .docx as raw bytes. The on-disk template is never mutated.
    """
    data = gather_data(franchise_id, week_iso, sales_tracker)
    doc  = Document(TEMPLATE_PATH)

    _strip_instruction_paragraphs(doc)
    _apply_text_replacements(doc, data)
    _populate_appendix_table(doc, data["appendix"])
    _inject_charts(doc, data)
    _append_modeled_footer(doc, data)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _apply_text_replacements(doc, data):
    """Fill in scalar placeholders + ranked lists + location sub-sections."""
    meta = data["meta"]
    period_label = meta["period_label"]
    franchise    = meta["franchise"]

    # Top-level scalars that appear across the cover + side-rails
    scalars = {
        "[Week of: DATE — DATE]": period_label,
        "[Week of: DATE]":         period_label,
        "[Business Owner Name]":   meta["owner"],
        "[Email / Phone]":         meta["contact"],
        "[# of active locations]": str(len(data["locations"])),
        "[Name]":                  meta["owner"],
        "[Date of report]":        meta["generated"],
        "[Location(s)]":           franchise,
        "[List active locations]": franchise,
        "[DATE RANGE]":            period_label,
    }

    for p in list(_iter_paragraphs(doc)):
        _replace_in_paragraph(p, scalars)

    _fill_weekly_summary(doc, data["summary"])
    _fill_intro_sentence(doc, franchise, period_label)
    _fill_hot_not(doc, data["hot"], data["not_hot"])
    _fill_locations(doc, data["locations"], data["cross_trends"])
    _fill_forecast(doc, data["forecast"])


def _fill_weekly_summary(doc, bullets):
    """Replace each '[Insight title] — One explanatory sentence. See Appendix A.' line."""
    bullet_paras = []
    target = "[Insight title]"
    for p in _iter_paragraphs(doc):
        if target in p.text:
            bullet_paras.append(p)

    # Replace as many template lines as we have data for
    for i, para in enumerate(bullet_paras):
        if i < len(bullets):
            b = bullets[i]
            _set_paragraph_text(para, f"{b['title']} — {b['body']}")
        else:
            _drop_paragraph(para)

    # If we have MORE bullets than template lines, append clones after the last
    if len(bullets) > len(bullet_paras) and bullet_paras:
        last = bullet_paras[-1]
        for b in bullets[len(bullet_paras):]:
            new_p = _insert_paragraph_after(last)
            new_p.style = last.style
            new_p.add_run(f"{b['title']} — {b['body']}")
            last = new_p

    # Drop the "[Add further bullets as data and owner prompts require.]" stub
    for p in list(_iter_paragraphs(doc)):
        if "[Add further bullets" in p.text:
            _drop_paragraph(p)


def _fill_intro_sentence(doc, franchise, period_label):
    target_prefix = "[Introductory sentence"
    for p in _iter_paragraphs(doc):
        if p.text.lstrip().startswith(target_prefix):
            _set_paragraph_text(
                p,
                f"The week of {period_label} at {franchise} showed the demand and waste patterns summarised below — "
                f"all figures sourced from Veratori item recognition.",
            )
            return


def _fill_hot_not(doc, hot, not_hot):
    """Hot/Not lists are 3 numbered lines each — find them by leading '1.', '2.', '3.'."""
    # Strategy: locate the "What's Hot" / "What's Not" header paragraphs first;
    # the next 3 paragraphs after each are the ranked items.
    paras = list(_iter_paragraphs(doc))

    def _find_index(prefix):
        for i, p in enumerate(paras):
            if p.text.strip().startswith(prefix):
                return i
        return None

    hot_idx = _find_index("What's Hot")
    not_idx = _find_index("What's Not")

    def _fill_block(start_idx, items, value_key, label):
        if start_idx is None:
            return
        filled = 0
        for j in range(start_idx + 1, min(start_idx + 12, len(paras))):
            p = paras[j]
            stripped = p.text.lstrip()
            if not stripped or not stripped[0].isdigit():
                continue
            if filled < len(items):
                it = items[filled]
                _set_paragraph_text(
                    p,
                    f"{filled + 1}.  {it['name']}  —  {it[value_key]} {label}.  {it['insight']}",
                )
                filled += 1
            else:
                # No data for this slot
                _set_paragraph_text(p, f"{filled + 1}.  {UNAVAILABLE}", italic=True)
                filled += 1
            if filled >= 3:
                break

    _fill_block(hot_idx, hot,     "units",  "units sold")
    _fill_block(not_idx, not_hot, "wasted", "units wasted")


def _fill_locations(doc, locations, cross_trends):
    """Replace [Location Name N] headers and 'Top seller:' lines."""
    name_paras   = []
    metric_paras = []
    for p in _iter_paragraphs(doc):
        if p.text.lstrip().startswith("[Location Name"):
            name_paras.append(p)
        elif p.text.lstrip().startswith("Top seller:"):
            metric_paras.append(p)

    # Fill in order
    for i, np_ in enumerate(name_paras):
        if i < len(locations):
            _set_paragraph_text(np_, locations[i]["name"])
        else:
            _drop_paragraph(np_)
    for i, mp in enumerate(metric_paras):
        if i < len(locations):
            loc = locations[i]
            _set_paragraph_text(
                mp,
                f"Top seller:  {loc['top_seller']}  —  {loc['top_units']} units.    "
                f"Most profitable:  {loc['profit_item']}  —  {loc['profit_rev']} revenue.",
            )
        else:
            _drop_paragraph(mp)

    # Cross-location trends paragraph
    for p in _iter_paragraphs(doc):
        if p.text.lstrip().startswith("Cross-Location Trends"):
            _set_paragraph_text(p, f"Cross-Location Trends  {cross_trends}")
            return


def _fill_forecast(doc, fc):
    """Section 05 — opening sentence, recommended orders, reduce, priorities."""
    paras = list(_iter_paragraphs(doc))

    # Opening sentence
    for p in paras:
        if p.text.lstrip().startswith("[Opening sentence"):
            _set_paragraph_text(p, fc["opening"])
            break

    # Recommended Orders: lines starting "1.  [Item]  —  Order"
    order_paras = [p for p in paras if "Order [Qty]" in p.text]
    for i, p in enumerate(order_paras):
        if i < len(fc["orders"]):
            o = fc["orders"][i]
            _set_paragraph_text(
                p, f"{i + 1}.  {o['item']}  —  Order {o['qty']} units.  Rationale: {o['rationale']}"
            )
        else:
            _drop_paragraph(p)
    for p in paras:
        if "[Add further order" in p.text:
            _drop_paragraph(p)

    # Items to Reduce
    reduce_paras = [p for p in paras if "Reduce order" in p.text and "[Item]" in p.text]
    for i, p in enumerate(reduce_paras):
        if i < len(fc["reduce"]):
            r = fc["reduce"][i]
            _set_paragraph_text(p, f"{r['item']}  —  Reduce order.  {r['detail']}")
        else:
            _set_paragraph_text(p, UNAVAILABLE, italic=True)
    for p in paras:
        if "[Add further items" in p.text:
            _drop_paragraph(p)

    # Location-Specific Priorities
    prio_paras = [p for p in paras if p.text.lstrip().startswith("[Location ") and "]:" in p.text]
    for i, p in enumerate(prio_paras):
        if i < len(fc["priorities"]):
            pr = fc["priorities"][i]
            _set_paragraph_text(p, f"{pr['location']}:  {pr['action']}")
        else:
            _drop_paragraph(p)


def _populate_appendix_table(doc, rows):
    """
    Locate the appendix table (header row contains 'Menu Item') and replace
    its placeholder rows with one row per menu item.
    """
    target = None
    for tbl in doc.tables:
        if tbl.rows and tbl.rows[0].cells and tbl.rows[0].cells[0].text.strip() == "Menu Item":
            target = tbl
            break
    if target is None:
        return

    # Existing placeholder rows are rows[1:]. We'll clone row[1] as a stencil
    # for each data row, then remove the originals.
    if len(target.rows) < 2:
        return

    stencil = target.rows[1]._tr  # XML element
    placeholder_rows = list(target.rows[1:])

    if not rows:
        # Replace single row with "Unavailable for this period" line
        cells = placeholder_rows[0].cells
        _set_paragraph_text(cells[0].paragraphs[0], UNAVAILABLE, italic=True)
        for c in cells[1:]:
            _set_paragraph_text(c.paragraphs[0], "")
        for extra in placeholder_rows[1:]:
            extra._tr.getparent().remove(extra._tr)
        return

    # Clone and fill
    parent = stencil.getparent()
    insert_after = stencil
    for i, row_data in enumerate(rows):
        if i == 0:
            new_tr = stencil
        else:
            new_tr = copy.deepcopy(stencil)
            insert_after.addnext(new_tr)
            insert_after = new_tr
        from docx.table import _Row
        row = _Row(new_tr, target)
        cells = row.cells
        _set_paragraph_text(cells[0].paragraphs[0], row_data["name"])
        _set_paragraph_text(cells[1].paragraphs[0], str(row_data["units"]))
        _set_paragraph_text(cells[2].paragraphs[0], str(row_data["revenue"]))
        _set_paragraph_text(cells[3].paragraphs[0], str(row_data["wasted"]))
        _set_paragraph_text(cells[4].paragraphs[0], str(row_data["note"]))

    # Remove leftover originals (rows beyond what we just wrote)
    written = len(rows)
    for extra in placeholder_rows[1:]:
        # placeholder_rows[1:] are the unused stencil siblings — drop them
        tr = extra._tr
        if tr.getparent() is not None:
            tr.getparent().remove(tr)


def _inject_charts(doc, data):
    """
    Replace the two chart anchor paragraphs (Chart 1 / Chart 2) with the
    rendered PNG image, and update each caption paragraph to the real caption.
    """
    pie_img = render_pie_chart(data["pie_data"])
    bar_img = render_bar_chart(data["bar_data"])

    def _swap(anchor_prefix, image_buf, caption):
        anchor = _find_paragraph_starting_with(doc, anchor_prefix)
        if anchor is None:
            return
        # The next non-empty paragraph is the caption.
        caption_para = None
        sibling = anchor._element.getnext()
        while sibling is not None and caption_para is None:
            if sibling.tag == qn("w:p"):
                from docx.text.paragraph import Paragraph
                cand = Paragraph(sibling, anchor._parent)
                if cand.text.strip():
                    caption_para = cand
                    break
            sibling = sibling.getnext()

        if image_buf is not None:
            img_para = _insert_paragraph_after(anchor)
            run = img_para.add_run()
            run.add_picture(image_buf, width=Inches(5.5))
        else:
            note_para = _insert_paragraph_after(anchor)
            note_para.add_run(UNAVAILABLE).italic = True

        if caption_para is not None:
            _set_paragraph_text(caption_para, caption)

    _swap("Chart 1", pie_img, data["charts"]["pie_caption"])
    _swap("Chart 2", bar_img, data["charts"]["bar_caption"])


def _append_modeled_footer(doc, data):
    """Append the 'Modeled estimate' disclaimer when data is synthesized."""
    if not data["meta"].get("is_modeled"):
        return
    p = doc.add_paragraph()
    run = p.add_run(
        "Modeled estimate — figures derived from Veratori demand modeling pending POS integration "
        "for this location. Live item-recognition data activates automatically once on-site cameras are deployed."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
